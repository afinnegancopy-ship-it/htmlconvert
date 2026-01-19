import streamlit as st
from docx import Document
from openpyxl import Workbook, load_workbook
from bs4 import BeautifulSoup
from datetime import datetime
from io import BytesIO
import re
import html
from docx.oxml.ns import qn

# =====================================================
# PAGE CONFIG
# =====================================================

st.set_page_config(
    page_title="Document Converter",
    page_icon="📄",
    layout="centered"
)

# =====================================================
# CUSTOM CSS - White Background, Black Text
# =====================================================

st.markdown("""
<style>
    .stApp {
        background-color: #ffffff;
        color: #000000;
    }
    
    .main-header {
        text-align: center;
        padding: 2rem 0 1rem 0;
    }
    
    .main-header h1 {
        color: #000000;
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    
    .main-header p {
        color: #000000;
        font-size: 1.1rem;
    }
    
    h1, h2, h3, h4, h5, h6, p, span, div, label {
        color: #000000 !important;
    }
    
    .stMarkdown {
        color: #000000;
    }
    
    .stRadio > div {
        display: flex;
        justify-content: center;
        gap: 1rem;
    }
    
    .stRadio > div > label {
        background-color: #f5f5f5;
        padding: 0.75rem 1.5rem;
        border-radius: 12px;
        border: 1px solid #e0e0e0;
        cursor: pointer;
        transition: all 0.2s;
        color: #000000 !important;
    }
    
    .stRadio > div > label:hover {
        border-color: #6366f1;
    }
    
    .stDownloadButton > button {
        background-color: #6366f1 !important;
        color: #ffffff !important;
        border: none !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        border-radius: 12px !important;
        width: 100% !important;
    }
    
    .stDownloadButton > button:hover {
        background-color: #818cf8 !important;
    }
    
    .stButton > button {
        color: #ffffff !important;
    }
    
    .success-box {
        background-color: #f0fdf4;
        border: 1px solid #22c55e;
        border-radius: 12px;
        padding: 1rem;
        text-align: center;
        margin: 1rem 0;
    }
    
    .success-box p {
        color: #16a34a !important;
        margin: 0;
    }
    
    div[data-testid="stFileUploader"] {
        background-color: #f9f9f9;
        border: 2px dashed #cccccc;
        border-radius: 16px;
        padding: 1rem;
    }
    
    div[data-testid="stFileUploader"]:hover {
        border-color: #6366f1;
    }
    
    div[data-testid="stFileUploader"] label {
        color: #000000 !important;
    }
    
    .stAlert {
        color: #000000;
    }
    
    footer {
        color: #000000 !important;
    }
</style>
""", unsafe_allow_html=True)

# =====================================================
# SPACING NORMALIZATION
# =====================================================

def normalize_spacing(text: str, preserve_edges: bool = False) -> str:
    if not text:
        return text
    
    text = text.replace('_x000D_', ' ')
    text = text.replace('_x000A_', ' ')
    
    leading_space = ""
    trailing_space = ""
    if preserve_edges:
        leading_match = re.match(r'^(\s+)', text)
        trailing_match = re.search(r'(\s+)$', text)
        if leading_match:
            leading_space = " "
        if trailing_match:
            trailing_space = " "
    
    text = html.unescape(text)
    text = text.replace("\xa0", " ")
    text = text.replace("\u00ad", "")
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = text.strip()
    
    if preserve_edges:
        return leading_space + text + trailing_space
    return text

# =====================================================
# WORD → HTML → EXCEL
# =====================================================

def run_is_bold(run):
    if run.bold is True:
        return True
    if run.bold is None:
        rPr = run._element.rPr
        if rPr is not None and rPr.find(qn('w:b')) is not None:
            return True
    return False

def paragraph_is_bold(paragraph):
    if paragraph.style and paragraph.style.font.bold:
        return True
    return False

def is_bullet_paragraph(paragraph):
    style_name = paragraph.style.name.lower()
    if 'list' in style_name or 'bullet' in style_name or 'number' in style_name:
        return True
    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        return True
    return False

def paragraph_to_html(paragraph):
    html_out = ""
    text = paragraph.text.strip()
    manual_bullet_match = re.match(r'^[\u2022\u00B7\-]\s+(.*)', text)
    is_bullet = is_bullet_paragraph(paragraph) or manual_bullet_match
    html_out += "<li>" if is_bullet else "<p>"
    strong_phrases = [
        "Description:", "How to Use:", "How To Use:", "Set Contains:",
        "Key Notes:", "Fit & Fabric", "Product Details",
        "Key Benefits", "Designed for"
    ]
    for run in paragraph.runs:
        run_text = normalize_spacing(run.text)
        if run_is_bold(run) or paragraph_is_bold(paragraph):
            run_text = f"<b>{run_text}</b>"
        for phrase in strong_phrases:
            run_text = run_text.replace(phrase, f"<strong>{phrase}</strong>")
        html_out += run_text
    html_out += "</li>" if is_bullet else "</p>"
    return html_out

def docx_to_html_blocks(docx_file):
    doc = Document(docx_file)
    html_blocks = {}
    current_id = None
    current_html = []
    inside_list = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.fullmatch(r'\d{8,}', text):
            if current_id and current_html:
                if inside_list:
                    current_html.append("</ul>")
                html_blocks[current_id] = normalize_spacing(''.join(current_html))
            current_id = text
            current_html = []
            inside_list = False
        else:
            is_bullet = is_bullet_paragraph(para)
            if is_bullet and not inside_list:
                current_html.append("<ul>")
                inside_list = True
            elif not is_bullet and inside_list:
                current_html.append("</ul>")
                inside_list = False
            current_html.append(paragraph_to_html(para))
    if current_id and current_html:
        if inside_list:
            current_html.append("</ul>")
        html_blocks[current_id] = normalize_spacing(''.join(current_html))
    return html_blocks

def export_html_to_excel(data):
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.append(["ID", "HTML"])
    for k, v in data.items():
        ws.append([k, v])
    wb.save(output)
    output.seek(0)
    return output

# =====================================================
# EXCEL (HTML) → WORD
# =====================================================

def add_inline_runs(paragraph, element):
    for child in element.children:
        tag_name = getattr(child, "name", None)
        
        if tag_name in ("b", "strong"):
            text = normalize_spacing(child.get_text())
            if text:
                run = paragraph.add_run(text)
                run.bold = True
        elif tag_name == "br":
            paragraph.add_run().add_break()
        elif tag_name == "li":
            p = paragraph._parent.add_paragraph(style="List Bullet")
            add_inline_runs(p, child)
        elif tag_name in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "span"):
            add_inline_runs(paragraph, child)
        elif tag_name is None:
            text = normalize_spacing(str(child), preserve_edges=True)
            if text:
                paragraph.add_run(text)
        else:
            add_inline_runs(paragraph, child)

def process_element(doc, element):
    tag_name = element.name
    nested_blocks = element.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'div'], recursive=False)
    
    if tag_name == "p" and nested_blocks:
        for child in element.children:
            if child.name in ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'div'):
                process_element(doc, child)
            elif child.name is None:
                text = normalize_spacing(str(child))
                if text:
                    doc.add_paragraph(text)
            elif child.name in ('b', 'strong'):
                p = doc.add_paragraph()
                run = p.add_run(normalize_spacing(child.get_text()))
                run.bold = True
            else:
                text = normalize_spacing(child.get_text())
                if text:
                    p = doc.add_paragraph()
                    add_inline_runs(p, child)
    elif tag_name == "p":
        p = doc.add_paragraph()
        add_inline_runs(p, element)
    elif tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        p = doc.add_paragraph()
        add_inline_runs(p, element)
    elif tag_name == "ul":
        for li in element.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, li)
    elif tag_name == "ol":
        for li in element.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, li)
    elif tag_name == "li":
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, element)
    elif tag_name == "div":
        for child in element.children:
            if hasattr(child, 'name') and child.name:
                process_element(doc, child)
            elif child.name is None:
                text = normalize_spacing(str(child))
                if text:
                    doc.add_paragraph(text)
    elif tag_name == "br":
        doc.add_paragraph()
    elif tag_name in ("strong", "b"):
        p = doc.add_paragraph()
        run = p.add_run(normalize_spacing(element.get_text()))
        run.bold = True
    else:
        text = normalize_spacing(element.get_text())
        if text:
            p = doc.add_paragraph()
            add_inline_runs(p, element)

def add_html_to_doc(doc, html_content):
    if not html_content:
        return
    
    html_content = str(html_content)
    html_content = html_content.replace('_x000D_', '')
    html_content = html_content.replace('_x000A_', '')
    
    if "<" not in html_content:
        doc.add_paragraph(normalize_spacing(html.unescape(html_content)))
        return
    
    html_content = html.unescape(html_content)
    soup = BeautifulSoup(html_content, "html.parser")
    
    for element in soup.contents:
        if element.name is None:
            text = normalize_spacing(str(element))
            if text:
                doc.add_paragraph(text)
        elif element.name:
            process_element(doc, element)

def excel_to_word(excel_file):
    wb = load_workbook(excel_file)
    ws = wb.active
    doc = Document()
    
    for row in ws.iter_rows(min_row=2, values_only=True):
        product_id, html_content = row
        if not product_id:
            continue
        p = doc.add_paragraph(str(product_id))
        p.runs[0].bold = True
        if html_content:
            add_html_to_doc(doc, html_content)
        doc.add_page_break()
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =====================================================
# MAIN APP
# =====================================================

# Header
st.markdown("""
<div class="main-header">
    <h1>📄 Document Converter</h1>
    <p>Word ⟷ HTML ⟷ Excel</p>
</div>
""", unsafe_allow_html=True)

# Mode Selection
st.markdown("### Select Conversion Mode")
mode = st.radio(
    "Conversion direction",
    options=["Word → Excel", "Excel → Word"],
    horizontal=True,
    label_visibility="collapsed"
)

st.markdown("---")

# File Upload
st.markdown("### Upload Your File")

if mode == "Word → Excel":
    uploaded_file = st.file_uploader(
        "Drop your Word file here or click to browse",
        type=["docx"],
        key="word_uploader"
    )
    output_extension = "xlsx"
    output_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
else:
    uploaded_file = st.file_uploader(
        "Drop your Excel file here or click to browse",
        type=["xlsx"],
        key="excel_uploader"
    )
    output_extension = "docx"
    output_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Convert Button & Download
if uploaded_file is not None:
    st.markdown("---")
    
    # Show file info
    st.markdown(f"**Selected file:** {uploaded_file.name}")
    
    if st.button("🔄 Convert", use_container_width=True, type="primary"):
        try:
            with st.spinner("Converting..."):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = uploaded_file.name.rsplit('.', 1)[0]
                output_filename = f"{base_name}_{timestamp}.{output_extension}"
                
                if mode == "Word → Excel":
                    data = docx_to_html_blocks(uploaded_file)
                    output_data = export_html_to_excel(data)
                else:
                    output_data = excel_to_word(uploaded_file)
            
            st.markdown('<div class="success-box"><p>✅ Conversion successful!</p></div>', unsafe_allow_html=True)
            
            st.download_button(
                label=f"📥 Download {output_extension.upper()}",
                data=output_data,
                file_name=output_filename,
                mime=output_mime,
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"❌ Conversion failed: {str(e)}")
else:
    st.info(f"👆 Please upload a {'Word (.docx)' if mode == 'Word → Excel' else 'Excel (.xlsx)'} file to begin.")

# Footer
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #000000; font-size: 0.85rem;'>Built with Streamlit</p>",
    unsafe_allow_html=True
)
